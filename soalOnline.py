import streamlit as st
from google import genai
from google.genai import types
from docx import Document
import io
import json

# --- 1. Konfigurasi Awal ---

# Ganti dengan Kunci API Google Gen AI Anda
GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]

# Inisialisasi Klien GenAI
try:
    client = genai.Client(api_key=GEMINI_API_KEY)
except Exception as e:
    # Hanya menampilkan error jika API Key belum disetel atau ada masalah
    client = None

# Nama Model yang akan digunakan
MODEL_NAME = "gemini-2.5-flash" 

# Peta Tingkat Kesukaran Soal ke Capaian Kognitif (C1-C6)
KESUKARAN_MAP = {
    "LOTS (C1-C2)": ["C1", "C2"],
    "MOTS (C3-C4)": ["C3", "C4"],
    "HOTS (C5-C6)": ["C5", "C6"]
}

# Mapping Fase ke Pilihan Kelas (Untuk memudahkan pengguna memilih)
FASE_KELAS_MAP = {
    "A": ["1", "2"],
    "B": ["3", "4"],
    "C": ["5", "6"],
    "D": ["7", "8", "9"],
    "E": ["10"],
    "F": ["11", "12"]
}

# --- 2. Antarmuka Streamlit (UI) ---

# Mengatur tata letak halaman menjadi lebar (wide layout)
st.set_page_config(
    page_title="Generator Soal & Kisi-Kisi Otomatis",
    layout="wide",
    initial_sidebar_state="collapsed" 
)

st.title("📝 Generator Soal & Kisi-Kisi Otomatis")
st.markdown("Aplikasi untuk membuat soal dan kisi-kisi berbasis kurikulum dan taksonomi Bloom.")

## ⚙️ Pengaturan Soal
st.header("1. Input Data dan Metadata Soal")
st.markdown("---")

col_metadata, col_counts = st.columns([2, 1])

with col_metadata:
    st.subheader("Data Kurikulum")
    
    # Metadata Kurikulum
    col_fase, col_kelas = st.columns(2)
    
    # 📌 PERUBAHAN UTAMA: Pilihan Fase dan Kelas
    fase = col_fase.selectbox("Fase Kurikulum", list(FASE_KELAS_MAP.keys()), index=4)
    
    # Dinamisasi pilihan kelas berdasarkan fase yang dipilih
    pilihan_kelas = FASE_KELAS_MAP.get(fase, [""])
    kelas = col_kelas.selectbox("Kelas", pilihan_kelas)
    
    mapel = st.text_input("Mata Pelajaran", "Kimia")
    materi = st.text_input("Materi Utama Soal", "Stoikiometri")
    sub_materi = st.text_area("Sub Materi Soal (Pisahkan dengan koma)", "Konsep Mol, Persamaan Reaksi Setara, Hukum Dasar Kimia")
    
    capaian_pembelajaran = st.text_area(
        "Definisi Capaian Pembelajaran (CP) untuk Soal ini", 
        "Peserta didik mampu menghitung jumlah mol, menentukan pereaksi pembatas, dan menerapkan konsep stoikiometri dalam perhitungan kimia sehari-hari."
    )

with col_counts:
    st.subheader("Format Soal & Jumlah")

    # Jenis Soal dan Jumlah
    pilihan_ganda_count = st.number_input("Pilihan Ganda (PG)", min_value=0, value=5)
    essay_count = st.number_input("Essay", min_value=0, value=2)
    isian_singkat_count = st.number_input("Isian Singkat", min_value=0, value=3)

    total_soal = pilihan_ganda_count + essay_count + isian_singkat_count
    st.info(f"Total Soal yang dibuat: **{total_soal}**")
    
    st.subheader("Distribusi Kesukaran")
    
    # Distribusi Tingkat Kesukaran
    lots_count = st.number_input(
        "LOTS (C1-C2)",
        min_value=0,
        max_value=total_soal,
        value=min(total_soal, 3) 
    )
    mots_count = st.number_input(
        "MOTS (C3-C4)",
        min_value=0,
        max_value=total_soal - lots_count,
        value=min(total_soal - lots_count, 4)
    )
    hots_count = st.number_input(
        "HOTS (C5-C6)",
        min_value=0,
        max_value=total_soal - lots_count - mots_count,
        value=total_soal - lots_count - mots_count
    )

    if (lots_count + mots_count + hots_count) != total_soal:
        st.error(f"Total Soal Kesukaran ({lots_count + mots_count + hots_count}) harus sama dengan Total Soal Dibuat ({total_soal}).")


# --- 3. Logika Pembuatan Soal dan Dokumen Word ---
# 📌 PERUBAHAN UTAMA: Menambahkan 'kelas' sebagai argumen di generate_questions dan create_word_document

def generate_questions(fase, kelas, mapel, materi, sub_materi, soal_counts, kesukaran_counts, cp, total_soal):
    """
    Membuat prompt dan memanggil Google Gen AI untuk menghasilkan soal.
    """
    if not client or total_soal == 0:
        return None, "Klien Gen AI tidak terinisialisasi atau total soal 0."

    # LOGIKA OPSI PG: Tentukan jumlah opsi berdasarkan Fase
    if fase in ["E", "F"]:
        num_options = 5  # Opsi A, B, C, D, E
    else:
        num_options = 4  # Opsi A, B, C, D
        
    opsi_list = ["A", "B", "C", "D"]
    if num_options == 5:
        opsi_list.append("E")
        
    opsi_fields = {opt: types.Schema(type=types.Type.STRING) for opt in opsi_list}
    
    # Distribusi C1-C6 (Logika ini tetap sama)
    c_levels_needed = []
    
    lots_c = KESUKARAN_MAP["LOTS (C1-C2)"]
    for i in range(kesukaran_counts.get("LOTS (C1-C2)", 0)):
        c_levels_needed.append(lots_c[i % len(lots_c)]) 
        
    mots_c = KESUKARAN_MAP["MOTS (C3-C4)"]
    for i in range(kesukaran_counts.get("MOTS (C3-C4)", 0)):
        c_levels_needed.append(mots_c[i % len(mots_c)])
        
    hots_c = KESUKARAN_MAP["HOTS (C5-C6)"]
    for i in range(kesukaran_counts.get("HOTS (C5-C6)", 0)):
        c_levels_needed.append(hots_c[i % len(hots_c)])

    if len(c_levels_needed) != total_soal:
        return None, "Kesalahan dalam distribusi C1-C6. Jumlah C-level tidak sesuai dengan total soal."

    question_specs = []
    current_index = 0
    pg_description = f"Pilihan Ganda ({num_options} opsi, 1 kunci jawaban)"
    
    for _ in range(soal_counts.get("PG", 0)):
        if current_index < len(c_levels_needed):
            question_specs.append({"type": pg_description, "kognitif": c_levels_needed[current_index]})
            current_index += 1
            
    for _ in range(soal_counts.get("Essay", 0)):
        if current_index < len(c_levels_needed):
            question_specs.append({"type": "Essay (Uraian)", "kognitif": c_levels_needed[current_index]})
            current_index += 1
            
    for _ in range(soal_counts.get("Isian Singkat", 0)):
        if current_index < len(c_levels_needed):
            question_specs.append({"type": "Isian Singkat (Jawaban 1 kata/frasa)", "kognitif": c_levels_needed[current_index]})
            current_index += 1

    formatted_specs = "\n".join([f"- Tipe: {spec['type']}, Kognitif: {spec['kognitif']}" for spec in question_specs])

    # Prompt Utama untuk Gemini (Menambahkan Kelas ke konteks prompt)
    system_prompt = f"""
    Anda adalah generator soal dan kisi-kisi profesional berbasis Taksonomi Bloom (C1-C6).
    Tugas Anda adalah menghasilkan soal dan kisi-kisi dalam format JSON yang spesifik.
    
    **Tujuan:** Membuat {total_soal} soal untuk {mapel} Kelas {kelas} (Fase {fase}).
    **Sub Materi:** {sub_materi}.
    **Capaian Pembelajaran (CP) Rujukan:** "{cp}"
    
    **Spesifikasi Soal:**
    {formatted_specs}
    
    **Instruksi Khusus:**
    1. Pastikan setiap soal unik dan sesuai dengan Tipe Soal dan Capaian Kognitif (C1-C6) yang diminta, serta **sesuai untuk tingkat Kelas {kelas}**.
    2. Untuk Pilihan Ganda, berikan {num_options} opsi ({', '.join(opsi_list)}) dan tentukan kunci jawabannya.
    3. Untuk semua tipe soal, sertakan Jawaban atau Kunci Jawaban.
    4. Buatlah indikator soal yang spesifik untuk setiap soal.
    5. JAWAB HANYA DALAM SATU BLOK JSON. JANGAN BERIKAN TEKS PENJELASAN APAPUN DI LUAR BLOK JSON.
    """

    # Struktur JSON yang Diharapkan (Tidak berubah)
    json_schema = types.Schema(
        type=types.Type.OBJECT,
        properties={
            "soal_dan_kisi_kisi": types.Schema(
                type=types.Type.ARRAY,
                description="Daftar semua soal, kisi-kisi, dan jawabannya.",
                items=types.Schema(
                    type=types.Type.OBJECT,
                    properties={
                        "nomor": types.Schema(type=types.Type.INTEGER, description="Nomor urut soal, dimulai dari 1."),
                        "tipe_soal": types.Schema(type=types.Type.STRING, enum=["Pilihan Ganda", "Essay", "Isian Singkat"]),
                        "cp_pembelajaran": types.Schema(type=types.Type.STRING, description=f"Capaian Pembelajaran yang sudah Disesuaikan dengan {mapel} dan disalin secara singkat."),
                        "indikator_soal": types.Schema(type=types.Type.STRING, description="Indikator soal spesifik dan terukur."),
                        "capaian_kognitif": types.Schema(type=types.Type.STRING, enum=["C1", "C2", "C3", "C4", "C5", "C6"]),
                        "pertanyaan": types.Schema(type=types.Type.STRING, description="Teks lengkap pertanyaan soal."),
                        "opsi_jawaban": types.Schema(
                            type=types.Type.OBJECT,
                            description=f"Hanya untuk Pilihan Ganda. Berisi {num_options} opsi (A sampai {opsi_list[-1]}).",
                            properties=opsi_fields 
                        ),
                        "kunci_jawaban": types.Schema(type=types.Type.STRING, description="Jawaban yang benar.")
                    },
                    required=["nomor", "tipe_soal", "cp_pembelajaran", "indikator_soal", "capaian_kognitif", "pertanyaan", "kunci_jawaban"]
                )
            )
        },
        required=["soal_dan_kisi_kisi"]
    )
    
    try:
        response = client.models.generate_content(
            model=MODEL_NAME,
            contents=system_prompt,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=json_schema,
            ),
        )
        result_json = json.loads(response.text)
        return result_json.get("soal_dan_kisi_kisi", []), None
        
    except Exception as e:
        return None, f"Terjadi kesalahan saat memanggil API: {e}"

def create_word_document(soal_data, mapel, materi, fase, kelas):
    """
    Membuat dokumen Word dari data soal dan kisi-kisi.
    """
    document = Document()
    
    # 1. Judul Dokumen (Menambahkan Kelas)
    document.add_heading(f"Soal {mapel} Kelas {kelas} - Fase {fase}", 0)
    document.add_paragraph(f"Materi: {materi}")
    document.add_paragraph("-" * 80)
    
    # 2. Daftar Soal
    document.add_heading("Daftar Soal", level=1)
    
    valid_opsi = ["A", "B", "C", "D"]
    if fase in ["E", "F"]:
        valid_opsi.append("E")
    
    for item in soal_data:
        tipe = item.get("tipe_soal", "")
        nomor = item.get("nomor", 0)
        
        if tipe == "Pilihan Ganda":
            document.add_heading(f"{nomor}. Pilihan Ganda", level=3)
            document.add_paragraph(item.get("pertanyaan", ""))
            
            opsi = item.get("opsi_jawaban", {})
            
            for key in valid_opsi:
                 if key in opsi and opsi[key] is not None:
                    val = opsi[key]
                    document.add_paragraph(f"  {key}. {val}")
                
            document.add_paragraph(f"  🔑 **Kunci Jawaban:** {item.get('kunci_jawaban', '')}")
            
        elif tipe == "Essay":
            document.add_heading(f"{nomor}. Essay", level=3)
            document.add_paragraph(item.get("pertanyaan", ""))
            document.add_paragraph(f"  📝 **Jawaban:** {item.get('kunci_jawaban', '')}")
            
        elif tipe == "Isian Singkat":
            document.add_heading(f"{nomor}. Isian Singkat", level=3)
            document.add_paragraph(item.get("pertanyaan", ""))
            document.add_paragraph(f"  📝 **Jawaban:** {item.get('kunci_jawaban', '')}")
            
        document.add_paragraph() 
        
    document.add_page_break()
    
    # 3. Kisi-Kisi Soal (Tidak berubah)
    document.add_heading("Kisi-Kisi Soal", level=1)
    
    table = document.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    
    hdr_cells[0].text = 'No.'
    hdr_cells[1].text = 'Nomor Soal'
    hdr_cells[2].text = 'Capaian Pembelajaran'
    hdr_cells[3].text = 'Indikator Soal'
    hdr_cells[4].text = 'Capaian Kognitif (C1-C6)'
    hdr_cells[5].text = 'Tipe Soal'
    
    for i, item in enumerate(soal_data):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = str(item.get('nomor', ''))
        row_cells[2].text = item.get('cp_pembelajaran', '')
        row_cells[3].text = item.get('indikator_soal', '')
        row_cells[4].text = item.get('capaian_kognitif', '')
        row_cells[5].text = item.get('tipe_soal', '')

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

# --- Logic Streamlit Utama (Tombol dan Hasil) ---

st.markdown("---")
st.header("2. Proses dan Hasil Generasi")

if st.button("🚀 Generate Soal dan Kisi-Kisi", type="primary", use_container_width=True):
    
    # Validasi
    if (lots_count + mots_count + hots_count) != total_soal:
        st.error("Gagal Generate: Total Soal Kesukaran (LOTS+MOTS+HOTS) harus sama dengan Total Soal Dibuat.")
    elif total_soal == 0:
         st.warning("Total soal adalah 0. Silakan atur jumlah soal.")
    elif not client:
        st.error("API Key belum disetel atau ada kesalahan inisialisasi. Mohon periksa GEMINI_API_KEY Anda.")
    else:
        with st.spinner("Membuat soal dan kisi-kisi menggunakan Google Gen AI..."):
            
            soal_counts = {
                "PG": pilihan_ganda_count,
                "Essay": essay_count,
                "Isian Singkat": isian_singkat_count
            }
            kesukaran_counts = {
                "LOTS (C1-C2)": lots_count,
                "MOTS (C3-C4)": mots_count,
                "HOTS (C5-C6)": hots_count
            }
            
            # Panggil fungsi generate (dengan 'kelas' sebagai argumen baru)
            soal_data, error_message = generate_questions(
                fase, kelas, mapel, materi, sub_materi, soal_counts, kesukaran_counts, capaian_pembelajaran, total_soal
            )
            
            if error_message:
                st.error(f"Gagal menghasilkan soal: {error_message}")
            elif soal_data:
                st.session_state.soal_data = soal_data
                st.success("✅ Soal dan Kisi-Kisi berhasil dibuat!")
                
                # Tampilkan hasil di UI (Kisi-Kisi)
                st.subheader("📊 Hasil Kisi-Kisi Soal")
                
                kisi_kisi_table = [
                    {
                        "No. Soal": item["nomor"],
                        "Tipe": item["tipe_soal"],
                        "CP Pembelajaran": item["cp_pembelajaran"],
                        "Indikator Soal": item["indikator_soal"],
                        "C. Kognitif": item["capaian_kognitif"]
                    } for item in soal_data
                ]
                st.dataframe(kisi_kisi_table, use_container_width=True)
                
                # Tampilkan hasil di UI (Soal)
                st.subheader("📖 Preview Soal")
                
                valid_opsi = ["A", "B", "C", "D"]
                if fase in ["E", "F"]:
                    valid_opsi.append("E")
                    
                for item in soal_data:
                    st.markdown(f"**{item['nomor']}. ({item['tipe_soal']} - {item['capaian_kognitif']})**")
                    st.write(item['pertanyaan'])
                    if item['tipe_soal'] == "Pilihan Ganda":
                        opsi = item.get('opsi_jawaban', {})
                        for key in valid_opsi:
                             if key in opsi and opsi[key] is not None:
                                 st.write(f"  - {key}. {opsi[key]}")

                    st.caption(f"Kunci Jawaban: **{item['kunci_jawaban']}**")
                    st.markdown("---")
            else:
                st.error("Gagal menghasilkan soal. Silakan coba lagi.")

# Tombol Download akan muncul setelah data soal ada di session state
if 'soal_data' in st.session_state and st.session_state.soal_data:
    st.subheader("📥 Unduh Hasil")
    
    # Panggil create_word_document dengan argumen 'kelas'
    word_file = create_word_document(
        st.session_state.soal_data, 
        mapel, 
        materi, 
        fase,
        kelas
    )
    
    st.download_button(
        label="Download sebagai Word (.docx)",
        data=word_file,
        file_name=f"Soal_{mapel}_Kelas_{kelas}_Fase_{fase}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )